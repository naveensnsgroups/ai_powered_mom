
import logging
from fpdf import FPDF
from docx import Document
from pydub import AudioSegment
import os
from typing import Dict, Any
import tempfile

logger = logging.getLogger(__name__)

# ---------------- Helper Functions ---------------- #

def save_temp_file(file_bytes: bytes, suffix: str = ".wav") -> str:
    """Save uploaded file temporarily and return the file path."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            return tmp.name
    except Exception as e:
        logger.error(f"Failed to save temp file: {str(e)}")
        raise

def validate_audio(file_bytes: bytes, min_size: int = 2000) -> bool:
    """Validate audio file size (or duration if needed)."""
    try:
        return len(file_bytes) >= min_size
    except Exception as e:
        logger.error(f"Audio validation failed: {str(e)}")
        return False

def convert_to_wav(file_path: str) -> str:
    """Convert audio to WAV format if not already WAV."""
    try:
        if file_path.endswith(".wav"):
            return file_path
        audio = AudioSegment.from_file(file_path)
        wav_path = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
        audio.export(wav_path, format="wav")
        return wav_path
    except Exception as e:
        logger.error(f"Failed to convert to WAV: {str(e)}")
        raise

# ---------------- Export Functions ---------------- #

def export_mom_pdf(mom: Dict[str, Any], output_path: str = "MoM.pdf") -> str:
    """Export structured MoM as PDF with all fields."""
    try:
        pdf = FPDF()
        pdf.add_page()

        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, mom.get("title", "Minutes of Meeting"), ln=True, align="C")
        pdf.set_font("Arial", '', 12)

        # Summary
        if "summary" in mom:
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Summary", ln=True)
            pdf.set_font("Arial", '', 12)
            pdf.multi_cell(0, 8, mom["summary"].get("summary", ""))
            pdf.multi_cell(0, 8, f"Overview: {mom.get('overview', '')}")

        # Attendees
        if mom.get("attendees"):
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Attendees", ln=True)
            pdf.set_font("Arial", '', 12)
            for attendee in mom["attendees"]:
                pdf.multi_cell(0, 8, f"- {attendee}")

        # Tasks
        if mom.get("tasks"):
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Tasks", ln=True)
            pdf.set_font("Arial", '', 12)
            for idx, item in enumerate(mom["tasks"], 1):
                pdf.multi_cell(
                    0, 8,
                    f"{idx}. Task: {item.get('task', '')} | Assigned to: {item.get('assigned_to', '')} | Deadline: {item.get('deadline', 'N/A')}"
                )

        # Action Items
        if mom.get("action_items"):
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Action Items", ln=True)
            pdf.set_font("Arial", '', 12)
            for idx, item in enumerate(mom["action_items"], 1):
                pdf.multi_cell(0, 8, f"{idx}. {item}")

        # Decisions
        if mom.get("decisions"):
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Decisions", ln=True)
            pdf.set_font("Arial", '', 12)
            for idx, item in enumerate(mom["decisions"], 1):
                pdf.multi_cell(
                    0, 8,
                    f"{idx}. Decision: {item.get('decision', '')} | Participant: {item.get('participant', '')}"
                )

        # Risks
        if mom.get("risks"):
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Risks", ln=True)
            pdf.set_font("Arial", '', 12)
            for idx, risk in enumerate(mom["risks"], 1):
                pdf.multi_cell(0, 8, f"{idx}. {risk}")

        # Data Points
        if mom.get("data_points"):
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Data Points", ln=True)
            pdf.set_font("Arial", '', 12)
            for idx, dp in enumerate(mom["data_points"], 1):
                pdf.multi_cell(0, 8, f"{idx}. {dp}")

        pdf.output(output_path)
        return os.path.abspath(output_path)

    except Exception as e:
        logger.error(f"Failed to export PDF: {str(e)}")
        raise

def export_mom_docx(mom: Dict[str, Any], output_path: str = "MoM.docx") -> str:
    """Export structured MoM as DOCX with all fields."""
    try:
        doc = Document()
        doc.add_heading(mom.get("title", "Minutes of Meeting"), 0)

        # Summary
        if "summary" in mom:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(mom["summary"].get("summary", ""))
            doc.add_paragraph(f"Overview: {mom.get('overview', '')}")

        # Attendees
        if mom.get("attendees"):
            doc.add_heading("Attendees", level=1)
            for attendee in mom["attendees"]:
                doc.add_paragraph(f"- {attendee}")

        # Tasks
        if mom.get("tasks"):
            doc.add_heading("Tasks", level=1)
            for idx, item in enumerate(mom["tasks"], 1):
                doc.add_paragraph(
                    f"{idx}. Task: {item.get('task', '')} | Assigned to: {item.get('assigned_to', '')} | Deadline: {item.get('deadline', 'N/A')}"
                )

        # Action Items
        if mom.get("action_items"):
            doc.add_heading("Action Items", level=1)
            for idx, item in enumerate(mom["action_items"], 1):
                doc.add_paragraph(f"{idx}. {item}")

        # Decisions
        if mom.get("decisions"):
            doc.add_heading("Decisions", level=1)
            for idx, item in enumerate(mom["decisions"], 1):
                doc.add_paragraph(
                    f"{idx}. Decision: {item.get('decision', '')} | Participant: {item.get('participant', '')}"
                )

        # Risks
        if mom.get("risks"):
            doc.add_heading("Risks", level=1)
            for idx, risk in enumerate(mom["risks"], 1):
                doc.add_paragraph(f"{idx}. {risk}")

        # Data Points
        if mom.get("data_points"):
            doc.add_heading("Data Points", level=1)
            for idx, dp in enumerate(mom["data_points"], 1):
                doc.add_paragraph(f"{idx}. {dp}")

        doc.save(output_path)
        return os.path.abspath(output_path)

    except Exception as e:
        logger.error(f"Failed to export DOCX: {str(e)}")
        raise
