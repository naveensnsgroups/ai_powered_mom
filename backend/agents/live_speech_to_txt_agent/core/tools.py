
import tempfile
import numpy as np
import soundfile as sf
import noisereduce as nr
from typing import Dict, List, Optional, Tuple
from agents.live_speech_to_txt_agent.core.agent import transcribe_chunk, run_live_agent
from pydub import AudioSegment
import io
import os
import logging
import subprocess
import wave
import struct

logger = logging.getLogger(__name__)

# ---------------- Helper Functions (unchanged from original) ---------------- #

def save_temp_file(file_bytes: bytes, suffix: str = ".wav") -> str:
    """Save uploaded audio temporarily and return the file path."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_bytes)
        return tmp.name

def validate_audio(file_bytes: bytes, min_size: int = 100) -> bool:
    """Validate audio file size (more lenient for chunks)."""
    return len(file_bytes) >= min_size

def is_valid_webm(file_bytes: bytes) -> bool:
    """Check if bytes contain a valid WebM file header."""
    if len(file_bytes) < 4:
        return False
    
    # WebM files start with EBML header (0x1A45DFA3)
    webm_signature = b'\x1a\x45\xdf\xa3'
    return file_bytes[:4] == webm_signature

def convert_raw_audio_to_wav(audio_bytes: bytes, sample_rate: int = 16000) -> Optional[str]:
    """
    Convert raw audio bytes to WAV format using multiple fallback methods.
    Returns WAV file path or None if conversion fails.
    """
    temp_paths_to_cleanup = []
    
    try:
        # Method 1: Try direct WAV creation if it's raw PCM data
        if len(audio_bytes) % 2 == 0:  # Assume 16-bit samples
            try:
                # Convert bytes to numpy array (assuming 16-bit PCM)
                audio_data = np.frombuffer(audio_bytes, dtype=np.int16).astype(np.float32) / 32768.0
                
                # Create WAV file
                wav_path = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
                sf.write(wav_path, audio_data, sample_rate)
                
                # Verify the file was created successfully
                if os.path.exists(wav_path) and os.path.getsize(wav_path) > 44:  # WAV header is 44 bytes
                    logger.info(f"Successfully converted raw PCM to WAV: {wav_path}")
                    return wav_path
                else:
                    os.remove(wav_path) if os.path.exists(wav_path) else None
            except Exception as e:
                logger.debug(f"Raw PCM conversion failed: {str(e)}")
        
        # Method 2: Save as temporary file and try pydub
        temp_input = save_temp_file(audio_bytes, suffix=".webm")
        temp_paths_to_cleanup.append(temp_input)
        
        try:
            audio_seg = AudioSegment.from_file(temp_input)
            audio_seg = audio_seg.set_channels(1).set_frame_rate(sample_rate)
            
            wav_path = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
            audio_seg.export(wav_path, format="wav")
            
            if os.path.exists(wav_path) and os.path.getsize(wav_path) > 44:
                logger.info(f"Successfully converted using pydub: {wav_path}")
                return wav_path
            else:
                os.remove(wav_path) if os.path.exists(wav_path) else None
                
        except Exception as e:
            logger.debug(f"pydub conversion failed: {str(e)}")
        
        # Method 3: Try FFmpeg with various input formats
        ffmpeg_formats = [
            ("webm", ["-f", "webm"]),
            ("ogg", ["-f", "ogg"]),
            ("raw", ["-f", "s16le", "-ar", str(sample_rate), "-ac", "1"]),
            ("auto", [])  # Let FFmpeg auto-detect
        ]
        
        for fmt_name, extra_args in ffmpeg_formats:
            try:
                wav_path = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
                
                cmd = ["ffmpeg", "-y"] + extra_args + [
                    "-i", temp_input,
                    "-ar", str(sample_rate),
                    "-ac", "1",
                    "-acodec", "pcm_s16le",
                    wav_path
                ]
                
                result = subprocess.run(
                    cmd, 
                    capture_output=True, 
                    text=True, 
                    timeout=30
                )
                
                if result.returncode == 0 and os.path.exists(wav_path) and os.path.getsize(wav_path) > 44:
                    logger.info(f"Successfully converted using FFmpeg ({fmt_name} format): {wav_path}")
                    return wav_path
                else:
                    os.remove(wav_path) if os.path.exists(wav_path) else None
                    logger.debug(f"FFmpeg {fmt_name} conversion failed: {result.stderr}")
                    
            except subprocess.TimeoutExpired:
                logger.warning(f"FFmpeg {fmt_name} conversion timed out")
                os.remove(wav_path) if os.path.exists(wav_path) else None
            except Exception as e:
                logger.debug(f"FFmpeg {fmt_name} conversion error: {str(e)}")
                if 'wav_path' in locals():
                    os.remove(wav_path) if os.path.exists(wav_path) else None
        
        # Method 4: Create silence if all else fails (for testing/debugging)
        logger.warning("All conversion methods failed, creating silence as fallback")
        return create_silence_wav(duration_ms=1000, sample_rate=sample_rate)
        
    except Exception as e:
        logger.error(f"Audio conversion completely failed: {str(e)}")
        return None
        
    finally:
        # Cleanup temporary files
        for path in temp_paths_to_cleanup:
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                except Exception as ex:
                    logger.warning(f"Failed to cleanup {path}: {str(ex)}")

def create_silence_wav(duration_ms: int = 1000, sample_rate: int = 16000) -> str:
    """Create a silent WAV file for fallback purposes."""
    samples = int(sample_rate * duration_ms / 1000)
    silence = np.zeros(samples, dtype=np.float32)
    
    wav_path = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
    sf.write(wav_path, silence, sample_rate)
    
    logger.info(f"Created silence WAV: {wav_path}")
    return wav_path

def reduce_noise_and_save(audio_bytes: bytes, sample_rate: int = 16000) -> str:
    """
    Convert audio bytes to clean WAV with noise reduction.
    Uses multiple fallback methods to handle various audio formats.
    """
    try:
        logger.debug(f"Processing audio chunk: {len(audio_bytes)} bytes")
        
        # Validate input
        if not validate_audio(audio_bytes, min_size=50):
            logger.warning(f"Audio chunk too small: {len(audio_bytes)} bytes")
            return create_silence_wav(duration_ms=500, sample_rate=sample_rate)
        
        # Convert to WAV
        wav_path = convert_raw_audio_to_wav(audio_bytes, sample_rate)
        if not wav_path:
            logger.warning("Audio conversion failed, using silence")
            return create_silence_wav(duration_ms=500, sample_rate=sample_rate)
        
        try:
            # Load and process audio
            data, sr = sf.read(wav_path, dtype="float32")
            
            # Handle stereo to mono conversion
            if data.ndim > 1:
                data = np.mean(data, axis=1)
            
            # Skip noise reduction if audio is too short
            if len(data) < sample_rate * 0.1:  # Less than 100ms
                logger.debug("Audio too short for noise reduction, skipping")
                reduced = data
            else:
                # Apply noise reduction with conservative settings
                try:
                    reduced = nr.reduce_noise(
                        y=data,
                        sr=sr,
                        stationary=False,  # Better for speech
                        prop_decrease=0.5,  # Conservative noise reduction
                        n_std_thresh_stationary=2.0  # Less aggressive
                    )
                except Exception as e:
                    logger.warning(f"Noise reduction failed: {str(e)}, using original audio")
                    reduced = data
            
            # Save final processed audio
            final_path = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
            sf.write(final_path, reduced, sr)
            
            # Cleanup intermediate file
            if os.path.exists(wav_path):
                os.remove(wav_path)
            
            logger.debug(f"Successfully processed audio: {final_path}")
            return final_path
            
        except Exception as e:
            logger.error(f"Audio processing failed: {str(e)}")
            # Cleanup and return original if processing fails
            if os.path.exists(wav_path):
                return wav_path  # Return original WAV
            else:
                return create_silence_wav(duration_ms=500, sample_rate=sample_rate)
                
    except Exception as e:
        logger.error(f"Audio preprocessing completely failed: {str(e)}")
        return create_silence_wav(duration_ms=500, sample_rate=sample_rate)

def load_audio_safe(file_path: str) -> Tuple[np.ndarray, int]:
    """
    Safely load audio file with multiple fallback methods.
    Returns (audio_data, sample_rate)
    """
    try:
        data, sr = sf.read(file_path, dtype="float32")
        
        # Convert stereo to mono if needed
        if data.ndim > 1:
            data = np.mean(data, axis=1)
            
        return data.astype(np.float32), int(sr)
        
    except Exception as e:
        logger.error(f"Failed to load audio {file_path}: {str(e)}")
        # Return silence as fallback
        return np.zeros(8000, dtype=np.float32), 16000

# ---------------- Enhanced LangChain Tools ---------------- #

def enhanced_live_transcribe_tool(audio_chunk: np.ndarray, sample_rate: int = 16000) -> Dict[str, str]:
    """
    Enhanced tool wrapper: live transcribes a single audio chunk.
    Returns transcript or error.
    """
    try:
        if len(audio_chunk) == 0:
            return {"transcript": "", "error": "Empty audio chunk"}
            
        result = transcribe_chunk(audio_chunk, sample_rate)
        return {"transcript": result or "", "error": ""}
    except Exception as e:
        logger.error(f"Error during enhanced live transcription: {str(e)}")
        return {"transcript": "", "error": f"Error during live transcription: {str(e)}"}

def enhanced_live_mom_tool(audio_chunks: List[np.ndarray], sample_rate: int = 16000) -> Dict[str, any]:
    """
    Enhanced tool wrapper: processes all chunks and generates comprehensive MoM.
    Returns enhanced MoM dict or error.
    """
    try:
        # Filter out empty chunks
        valid_chunks = [chunk for chunk in audio_chunks if len(chunk) > 0]
        
        if not valid_chunks:
            return {
                "mom": create_enhanced_fallback_response(),
                "error": ""
            }
            
        result = run_live_agent(valid_chunks, sample_rate)
        return {"mom": result, "error": ""}
    except Exception as e:
        logger.error(f"Error generating enhanced live MoM: {str(e)}")
        return {
            "mom": create_enhanced_fallback_response(f"Error: {str(e)}"),
            "error": f"Error generating live MoM: {str(e)}"
        }

def create_enhanced_fallback_response(error_msg: str = "No valid audio chunks received") -> Dict[str, any]:
    """Create fallback response with enhanced MoM structure."""
    from datetime import datetime
    
    return {
        "transcript": "",
        "mom": {
            "meeting_info": {
                "date": datetime.now().strftime("%Y-%m-%d"),
                "time": "Not specified",
                "meeting_type": "general meeting"
            },
            "attendance": {
                "participants": [],
                "total_participants": 0
            },
            "summary": {
                "overview": "Processing failed",
                "detailed": error_msg,
                "key_topics": []
            },
            "action_items": [],
            "decisions": [],
            "follow_up": {
                "next_meeting": "TBD",
                "pending_items": [],
                "required_approvals": []
            },
            "risks_and_blockers": []
        }
    }

# ---------------- Export Functions for Enhanced MoM ---------------- #

def export_enhanced_mom_pdf(mom_data: Dict[str, any], output_path: str = None) -> str:
    """
    Export enhanced MoM to PDF format with comprehensive layout.
    """
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
        import tempfile
        from datetime import datetime

        if output_path is None:
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name

        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=A4, 
                              rightMargin=72, leftMargin=72, 
                              topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkblue,
            borderWidth=1,
            borderColor=colors.darkblue,
            borderRadius=5,
            backColor=colors.lightblue,
            borderPadding=8
        )
        
        # Build story
        story = []
        
        # Title
        story.append(Paragraph("Meeting Minutes", title_style))
        story.append(Spacer(1, 20))
        
        # Meeting Info
        meeting_info = mom_data.get("mom", {}).get("meeting_info", {})
        info_data = [
            ["Date:", meeting_info.get("date", "Not specified")],
            ["Time/Duration:", meeting_info.get("time", "Not specified")],
            ["Meeting Type:", meeting_info.get("meeting_type", "General Meeting").title()],
            ["Generated:", datetime.now().strftime("%Y-%m-%d %H:%M:%S")]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Attendance
        story.append(Paragraph("Attendance", heading_style))
        attendance = mom_data.get("mom", {}).get("attendance", {})
        participants = attendance.get("participants", [])
        
        if participants:
            attendance_data = [["Name", "Role", "Status"]]
            for participant in participants:
                attendance_data.append([
                    participant.get("name", "Unknown"),
                    participant.get("role", "Not specified"),
                    participant.get("attendance_status", "Present").title()
                ])
            
            attendance_table = Table(attendance_data, colWidths=[2*inch, 2*inch, 1.5*inch])
            attendance_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(attendance_table)
        else:
            story.append(Paragraph("No participants identified", styles['Normal']))
        
        story.append(Spacer(1, 15))
        
        # Summary
        story.append(Paragraph("Meeting Summary", heading_style))
        summary = mom_data.get("mom", {}).get("summary", {})
        
        story.append(Paragraph("<b>Overview:</b>", styles['Normal']))
        story.append(Paragraph(summary.get("overview", "No overview available"), styles['Normal']))
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("<b>Detailed Notes:</b>", styles['Normal']))
        story.append(Paragraph(summary.get("detailed", "No detailed notes available"), styles['Normal']))
        
        # Key Topics
        key_topics = summary.get("key_topics", [])
        if key_topics:
            story.append(Spacer(1, 10))
            story.append(Paragraph("<b>Key Topics:</b>", styles['Normal']))
            for topic in key_topics:
                story.append(Paragraph(f"• {topic}", styles['Normal']))
        
        story.append(Spacer(1, 15))
        
        # Action Items
        story.append(Paragraph("Action Items", heading_style))
        action_items = mom_data.get("mom", {}).get("action_items", [])
        
        if action_items:
            action_data = [["#", "Task", "Assigned To", "Deadline", "Priority"]]
            for item in action_items:
                action_data.append([
                    str(item.get("id", "")),
                    item.get("task", "No description")[:50] + ("..." if len(item.get("task", "")) > 50 else ""),
                    item.get("assigned_to", "Not assigned"),
                    item.get("deadline", "No deadline"),
                    item.get("priority", "Medium").title()
                ])
            
            action_table = Table(action_data, colWidths=[0.5*inch, 2.5*inch, 1.5*inch, 1*inch, 0.8*inch])
            action_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            story.append(action_table)
        else:
            story.append(Paragraph("No action items identified", styles['Normal']))
        
        story.append(Spacer(1, 15))
        
        # Decisions
        story.append(Paragraph("Decisions", heading_style))
        decisions = mom_data.get("mom", {}).get("decisions", [])
        
        if decisions:
            for i, decision in enumerate(decisions, 1):
                story.append(Paragraph(f"<b>Decision {i}:</b> {decision.get('decision', 'No description')}", styles['Normal']))
                if decision.get("rationale"):
                    story.append(Paragraph(f"<i>Rationale:</i> {decision.get('rationale')}", styles['Normal']))
                if decision.get("responsible_party"):
                    story.append(Paragraph(f"<i>Responsible:</i> {decision.get('responsible_party')}", styles['Normal']))
                story.append(Spacer(1, 8))
        else:
            story.append(Paragraph("No decisions recorded", styles['Normal']))
        
        story.append(Spacer(1, 15))
        
        # Follow-up
        story.append(Paragraph("Follow-up", heading_style))
        follow_up = mom_data.get("mom", {}).get("follow_up", {})
        
        story.append(Paragraph(f"<b>Next Meeting:</b> {follow_up.get('next_meeting', 'TBD')}", styles['Normal']))
        
        pending_items = follow_up.get("pending_items", [])
        if pending_items:
            story.append(Paragraph("<b>Pending Items:</b>", styles['Normal']))
            for item in pending_items:
                story.append(Paragraph(f"• {item}", styles['Normal']))
        
        # Risks and Blockers
        risks = mom_data.get("mom", {}).get("risks_and_blockers", [])
        if risks:
            story.append(Spacer(1, 15))
            story.append(Paragraph("Risks & Blockers", heading_style))
            for risk in risks:
                severity = risk.get("severity", "Unknown").upper()
                story.append(Paragraph(f"<b>[{severity}]</b> {risk.get('issue', 'No description')}", styles['Normal']))
                if risk.get("owner"):
                    story.append(Paragraph(f"<i>Owner:</i> {risk.get('owner')}", styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"Enhanced MoM exported to PDF: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Failed to export enhanced MoM to PDF: {str(e)}")
        # Create simple fallback PDF
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            
            doc = SimpleDocTemplate(output_path or tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name)
            styles = getSampleStyleSheet()
            story = [Paragraph(f"Export Error: {str(e)}", styles['Normal'])]
            doc.build(story)
            return output_path
        except:
            raise e

def export_enhanced_mom_docx(mom_data: Dict[str, any], output_path: str = None) -> str:
    """
    Export enhanced MoM to DOCX format with comprehensive layout.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        import tempfile
        from datetime import datetime

        if output_path is None:
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name

        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('Meeting Minutes', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Meeting Info
        meeting_info = mom_data.get("mom", {}).get("meeting_info", {})
        doc.add_heading('Meeting Information', level=1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Date:', meeting_info.get("date", "Not specified")),
            ('Time/Duration:', meeting_info.get("time", "Not specified")),
            ('Meeting Type:', meeting_info.get("meeting_type", "General Meeting").title()),
            ('Generated:', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
        
        # Attendance
        doc.add_heading('Attendance', level=1)
        attendance = mom_data.get("mom", {}).get("attendance", {})
        participants = attendance.get("participants", [])
        
        if participants:
            attendance_table = doc.add_table(rows=1, cols=3)
            attendance_table.style = 'Table Grid'
            
            # Header
            hdr_cells = attendance_table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Role'
            hdr_cells[2].text = 'Status'
            
            # Add participants
            for participant in participants:
                row_cells = attendance_table.add_row().cells
                row_cells[0].text = participant.get("name", "Unknown")
                row_cells[1].text = participant.get("role", "Not specified")
                row_cells[2].text = participant.get("attendance_status", "Present").title()
        else:
            doc.add_paragraph("No participants identified")
        
        # Summary
        doc.add_heading('Meeting Summary', level=1)
        summary = mom_data.get("mom", {}).get("summary", {})
        
        doc.add_paragraph().add_run('Overview: ').bold = True
        doc.add_paragraph(summary.get("overview", "No overview available"))
        
        doc.add_paragraph().add_run('Detailed Notes: ').bold = True
        doc.add_paragraph(summary.get("detailed", "No detailed notes available"))
        
        # Key Topics
        key_topics = summary.get("key_topics", [])
        if key_topics:
            doc.add_paragraph().add_run('Key Topics: ').bold = True
            for topic in key_topics:
                doc.add_paragraph(f'• {topic}')
        
        # Action Items
        doc.add_heading('Action Items', level=1)
        action_items = mom_data.get("mom", {}).get("action_items", [])
        
        if action_items:
            for item in action_items:
                p = doc.add_paragraph()
                p.add_run(f"#{item.get('id', '')}: ").bold = True
                p.add_run(item.get("task", "No description"))
                
                details = []
                if item.get("assigned_to") and item.get("assigned_to") != "Not specified":
                    details.append(f"Assigned to: {item.get('assigned_to')}")
                if item.get("deadline") and item.get("deadline") not in ["N/A", "Not specified"]:
                    details.append(f"Deadline: {item.get('deadline')}")
                if item.get("priority"):
                    details.append(f"Priority: {item.get('priority').title()}")
                
                if details:
                    detail_p = doc.add_paragraph()
                    detail_p.add_run(" | ".join(details)).italic = True
        else:
            doc.add_paragraph("No action items identified")
        
        # Decisions
        doc.add_heading('Decisions', level=1)
        decisions = mom_data.get("mom", {}).get("decisions", [])
        
        if decisions:
            for i, decision in enumerate(decisions, 1):
                p = doc.add_paragraph()
                p.add_run(f"Decision {i}: ").bold = True
                p.add_run(decision.get("decision", "No description"))
                
                if decision.get("rationale"):
                    rat_p = doc.add_paragraph()
                    rat_p.add_run("Rationale: ").italic = True
                    rat_p.add_run(decision.get("rationale"))
                
                if decision.get("responsible_party"):
                    resp_p = doc.add_paragraph()
                    resp_p.add_run("Responsible: ").italic = True
                    resp_p.add_run(decision.get("responsible_party"))
        else:
            doc.add_paragraph("No decisions recorded")
        
        # Follow-up
        doc.add_heading('Follow-up', level=1)
        follow_up = mom_data.get("mom", {}).get("follow_up", {})
        
        doc.add_paragraph().add_run(f"Next Meeting: {follow_up.get('next_meeting', 'TBD')}")
        
        pending_items = follow_up.get("pending_items", [])
        if pending_items:
            doc.add_paragraph().add_run('Pending Items:').bold = True
            for item in pending_items:
                doc.add_paragraph(f"• {item}")
        
        # Risks and Blockers
        risks = mom_data.get("mom", {}).get("risks_and_blockers", [])
        if risks:
            doc.add_heading('Risks & Blockers', level=1)
            for risk in risks:
                p = doc.add_paragraph()
                severity = risk.get("severity", "Unknown").upper()
                p.add_run(f"[{severity}] ").bold = True
                p.add_run(risk.get("issue", "No description"))
                
                if risk.get("owner"):
                    owner_p = doc.add_paragraph()
                    owner_p.add_run("Owner: ").italic = True
                    owner_p.add_run(risk.get("owner"))
        
        # Save document
        doc.save(output_path)
        
        logger.info(f"Enhanced MoM exported to DOCX: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Failed to export enhanced MoM to DOCX: {str(e)}")
        # Create simple fallback DOCX
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('Export Error', 0)
            doc.add_paragraph(f"Failed to export meeting minutes: {str(e)}")
            doc.save(output_path or tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name)
            return output_path
        except:
            raise e

# ---------------- Legacy compatibility functions ---------------- #

def live_transcribe_tool(audio_chunk: np.ndarray, sample_rate: int = 16000) -> Dict[str, str]:
    """Legacy wrapper for enhanced live transcribe tool."""
    return enhanced_live_transcribe_tool(audio_chunk, sample_rate)

def live_mom_tool(audio_chunks: List[np.ndarray], sample_rate: int = 16000) -> Dict[str, str]:
    """Legacy wrapper for enhanced live MoM tool."""
    result = enhanced_live_mom_tool(audio_chunks, sample_rate)
    # Convert back to legacy format if needed
    if "mom" in result and isinstance(result["mom"], dict) and "mom" in result["mom"]:
        # Extract the nested mom structure for backward compatibility
        result["mom"] = result["mom"]["mom"]
    return result


